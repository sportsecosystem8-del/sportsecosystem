"""
Generate professional System Design & Panel Architecture Q&A document.
Output: docs/Sports_Ecosystem_System_Design_QA.docx
Run: python docs/build_system_design_qa_docx.py
"""
from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor


OUT = Path(__file__).resolve().parent / "Sports_Ecosystem_System_Design_QA.docx"


def set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_hr(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1B4F72")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_title_page(doc: Document) -> None:
    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SPORTS ECOSYSTEM PLATFORM")
    set_run_font(r, 22, True, RGBColor(0x1B, 0x4F, 0x72))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("System Design & Technical Architecture")
    set_run_font(r2, 16, True, RGBColor(0x2E, 0x86, 0xAB))

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Panel Flows · Security · Matching · Payments")
    set_run_font(r3, 12, False, RGBColor(0x55, 0x55, 0x55))

    doc.add_paragraph()
    add_hr(doc.add_paragraph())

    meta = [
        ("Project Group", "FYP26-CS-G22"),
        ("Document Type", "System Design Q&A / Technical Reference"),
        ("Stack", "Node.js · Express · MongoDB · React (Vite) · JWT"),
        ("Date", date.today().strftime("%d %B %Y")),
        ("Version", "1.0"),
    ]
    for label, value in meta:
        row = doc.add_paragraph()
        row.alignment = WD_ALIGN_PARAGRAPH.CENTER
        a = row.add_run(f"{label}: ")
        set_run_font(a, 11, True)
        b = row.add_run(value)
        set_run_font(b, 11)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(
        "This document describes how the implemented system is designed and how "
        "major features work across Admin, Player, Coach, and Business panels."
    )
    set_run_font(nr, 10, False, RGBColor(0x66, 0x66, 0x66))
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    items = [
        "1. System Design Overview",
        "2. Capacity, Scaling & User Handling",
        "3. Security & Privacy",
        "4. Role-Based Registration",
        "5. Authentication & Session Management",
        "6. Admin Panel",
        "7. Player Panel",
        "8. Coach Panel",
        "9. Business Panel",
        "10. Key Source Files Reference",
    ]
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
    doc.add_page_break()


def add_table(doc: Document, headers, rows) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = str(val)
    doc.add_paragraph()


def add_qa(doc: Document, q: str, a_paras: list) -> None:
    pq = doc.add_paragraph()
    rq = pq.add_run(q)
    set_run_font(rq, 11, True, RGBColor(0x1B, 0x4F, 0x72))
    for text in a_paras:
        pa = doc.add_paragraph(text)
        pa.paragraph_format.left_indent = Inches(0.15)


def build() -> Path:
    doc = Document()
    set_normal_style(doc)
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    add_title_page(doc)
    add_toc(doc)

    # --- 1 ---
    doc.add_heading("1. System Design Overview", level=1)
    doc.add_paragraph(
        "The Sports Ecosystem Platform is a role-based web application that connects "
        "players, coaches, and sports businesses. It supports coach discovery and matching, "
        "training lifecycle management, indoor ground booking, an equipment marketplace, "
        "subscriptions, and admin verification."
    )
    doc.add_paragraph("High-level architecture:", style="List Bullet")
    for line in (
        "Frontend: React (Vite) SPA with role-specific dashboards",
        "Backend: Node.js + Express REST API",
        "Database: MongoDB via Mongoose",
        "Auth: JWT (stateless) + bcrypt password hashing",
        "Deploy: API on Render (or similar); SPA on Vercel; MongoDB Atlas",
    ):
        doc.add_paragraph(line, style="List Bullet")

    # --- 2 ---
    doc.add_heading("2. Capacity, Scaling & User Handling", level=1)

    add_qa(
        doc,
        "Q1. How is the system designed for concurrent users?",
        [
            "There is no hard-coded concurrent-user cap. Capacity is achieved through "
            "process clustering, database connection pooling, and production rate limiting.",
        ],
    )
    add_table(
        doc,
        ["Mechanism", "Behaviour", "Location"],
        [
            ("PM2 cluster", "instances = max (CPU cores); 512MB memory restart", "ecosystem.config.js"),
            ("MongoDB pool", "maxPoolSize 50 · minPoolSize 5", "backend/config/database.js"),
            ("API rate limit", "5000 requests / 15 min (production)", "backend/middleware/rateLimit.js"),
            ("Auth rate limit", "150 requests / 15 min (production)", "backend/middleware/rateLimit.js"),
            ("Render default", "Single Node process (npm start)", "render.yaml"),
        ],
    )
    doc.add_paragraph(
        "Authenticated API calls are keyed by JWT subject (user id); auth routes are keyed by IP. "
        "Practical concurrency depends on host CPU, MongoDB Atlas tier, and rate-limit settings."
    )

    add_qa(
        doc,
        "Q2. What technologies handle users?",
        [
            "User accounts and profiles are stored in MongoDB (User + role-specific profile models). "
            "HTTP is served by Express. The frontend stores the JWT in localStorage and sends it on "
            "each API call via the Axios Authorization header.",
        ],
    )
    add_table(
        doc,
        ["Layer", "Technology"],
        [
            ("Runtime", "Node.js ≥ 18"),
            ("API", "Express 4"),
            ("Database", "MongoDB + Mongoose 8"),
            ("Auth tokens", "jsonwebtoken (JWT)"),
            ("Passwords", "bcryptjs (cost factor 12)"),
            ("Validation", "express-validator"),
            ("Uploads", "multer → backend/uploads/"),
            ("Frontend", "React + Vite + Axios"),
        ],
    )

    # --- 3 ---
    doc.add_heading("3. Security & Privacy", level=1)
    add_qa(
        doc,
        "Q3. What is used for system and user security/privacy?",
        [
            "Multiple layers protect authentication, authorization, uploads, and secrets:",
        ],
    )
    add_table(
        doc,
        ["Control", "Implementation"],
        [
            ("Password storage", "bcrypt hash, cost 12 — never plaintext"),
            ("Session / auth", "Signed JWT with user id (sub) and role"),
            ("Email / reset tokens", "crypto random token; only SHA-256 hash stored"),
            ("Strong passwords", "Min 8 chars; upper, lower, digit, special"),
            ("RBAC", "authenticate + requireRole middleware on routes"),
            ("Rate limiting", "Production API & auth limiters"),
            ("CORS", "CLIENT_URL origin in production"),
            ("Upload safety", "MIME filter, 8MB limit, sanitized filenames"),
            ("Sensitive docs", "Auth-protected stream (not public listing)"),
            ("Secrets", "Environment variables (.env) — not committed"),
            ("Account deletion", "Password + confirm DELETE; hard-delete frees email"),
        ],
    )
    doc.add_paragraph(
        "Key files: backend/middleware/auth.js, backend/middleware/rateLimit.js, "
        "backend/middleware/upload.js, backend/controllers/authController.js, "
        "backend/utils/deleteUserAccount.js, backend/app.js."
    )

    # --- 4 ---
    doc.add_heading("4. Role-Based Registration", level=1)
    add_qa(
        doc,
        "Q4. What is used for role-based registration, where is the code, and how does it work?",
        [
            "Primary handler: register() in backend/controllers/authController.js",
            "Route: POST /api/auth/register (backend/routes/auth.routes.js)",
            "UI: frontend/src/pages/Register.jsx",
        ],
    )
    doc.add_paragraph("Flow:", style="List Bullet")
    for step in (
        "Client sends email, password, role (player | coach | business_owner), and profile payload.",
        "Admin role cannot be registered via the public API (403).",
        "Role-specific validation runs (coach specialties, cricket category, business address/map).",
        "Password is hashed with bcrypt; User document is created.",
        "Matching profile is created: PlayerProfile, CoachProfile, or BusinessProfile.",
        "Players start as verificationStatus = verified; coaches/business start as pending_review.",
        "Email verification token is issued; user must verify before login (non-admin).",
        "No JWT is returned on register — login happens after email verification.",
    ):
        doc.add_paragraph(step, style="List Number")

    # --- 5 ---
    doc.add_heading("5. Authentication & Session Management", level=1)
    add_qa(
        doc,
        "Q5. What is used for authentication and how does it work?",
        [
            "JWT-based authentication (stateless — no server session store).",
        ],
    )
    doc.add_paragraph("Sequence:", style="List Bullet")
    for step in (
        "Login: POST /api/auth/login → credentials verified → signToken(user).",
        "JWT payload: { sub: userId, role } signed with JWT_SECRET.",
        "Frontend stores token in localStorage (useAuth.jsx).",
        "Axios attaches Authorization: Bearer <token> (services/api.js).",
        "authenticate middleware verifies JWT and sets req.user = { id, role }.",
        "requireRole(...) enforces RBAC on protected routes.",
        "GET /api/auth/me returns current user + populated profile.",
    ):
        doc.add_paragraph(step, style="List Number")

    add_qa(
        doc,
        "Q6. How long is one session?",
        [
            "Default JWT lifetime is 7 days (JWT_EXPIRES_IN=7d, with code fallback '7d' in signToken). "
            "There is no separate server-side session TTL. When the token expires, the user must log in again. "
            "A 401 response clears the token and redirects to login.",
        ],
    )

    # --- 6 Admin ---
    doc.add_heading("6. Admin Panel", level=1)

    add_qa(
        doc,
        "Q7. When an admin approves a user, how does it work?",
        [
            "Coaches and business owners require admin verification after email verification.",
            "Admin lists pending users (verificationStatus in pending_review / more_info).",
            "PATCH /api/admin/verification/coaches/:userId (or business) with action approve|reject|more_info.",
            "On approve: verificationStatus = verified; verification documents marked approved; "
            "in-app notification and optional approval email are sent.",
            "Code: patchCoachVerification / patchBusinessVerification in adminController.js; "
            "UI: AdminVerifyCoaches.jsx / AdminVerifyBusiness.jsx.",
        ],
    )

    add_qa(
        doc,
        "Q8. What is used to send documents, and what is the full flow to the admin?",
        [
            "Upload stack: multer middleware (PDF/JPG/PNG, max 8MB) writing to backend/uploads/.",
            "Coach: POST /api/coaches/documents → VerificationDocument (roleContext coach, status pending).",
            "Business: POST /api/business/documents → same model with business_owner context.",
            "Upload sets user verificationStatus to pending_review.",
            "Admin opens verification queue, previews via authenticated file stream "
            "(GET /admin/verification/documents/:docId/file), then approves the user or individual documents.",
            "Files: middleware/upload.js, models/VerificationDocument.js, coach/business controllers, "
            "utils/streamVerificationDocument.js.",
        ],
    )

    add_qa(
        doc,
        "Q9. If admin deletes a sport category, does it delete from existing users?",
        [
            "No. DELETE /api/admin/sports/:id only removes the SportCategory catalog document. "
            "Player and coach profiles store sport as string enums (e.g. cricket, badminton), not foreign keys. "
            "Existing profile data remains unchanged. There is no cascade delete to users.",
        ],
    )

    # --- 7 Player ---
    doc.add_heading("7. Player Panel", level=1)

    add_qa(
        doc,
        "Q10. On what basis are coaches shown to the player? What is the matching flow?",
        [
            "Endpoint: GET /api/players/recommendations (getRecommendations in playerController.js).",
            "Hard filters: verified, not suspended, matching sport; cricket also matches coaching category "
            "(batting / bowling / all-rounder).",
            "Baseline weighted score (then optional AI candidate selection; final order by baseline score):",
        ],
    )
    add_table(
        doc,
        ["Factor", "Weight", "Logic"],
        [
            ("Skill", "32%", "Sport specialty + preferredPlayerLevels + ratings/experience"),
            ("Schedule / time", "23%", "Shared preferred training days"),
            ("Location", "18%", "City match (exact / near / different)"),
            ("Category", "15%", "Cricket coaching focus vs player role"),
            ("Performance", "12%", "Recent evaluations vs coach profile signal"),
        ],
    )
    doc.add_paragraph(
        "AI (optional, via aiCoachEngine.js) may select candidates; ranking for display uses baseline "
        "matchScore. The UI intentionally does not show numeric scores to players — only ranked coaches."
    )
    doc.add_paragraph(
        "Supporting files: playerController.js, services/aiCoachEngine.js, utils/aiSchemas.js, "
        "frontend pages PlayerCoaches.jsx / PlayerDashboard.jsx."
    )

    add_qa(
        doc,
        "Q11. How is the player profile photo uploaded? Which function?",
        [
            "Route: POST /api/players/me/profile-photo with multer uploadImage.single('image').",
            "Controller: uploadProfilePhoto in playerController.js.",
            "Stores path '/uploads/<filename>' on PlayerProfile.profilePhotoUrl.",
            "Frontend: PlayerProfile.jsx (onPhoto) builds FormData and posts the image field.",
        ],
    )

    add_qa(
        doc,
        "Q12. How do images generally upload in the system?",
        [
            "Central middleware: backend/middleware/upload.js (upload for docs; uploadImage for photos).",
            "Files land on disk under backend/uploads/; Express serves them at /uploads.",
            "Callers store the relative path on the relevant model (profile, store logo, product images).",
            "Generic authenticated upload: POST /api/uploads/image (uploadController.js).",
            "Frontend resolves URLs with publicAssetUrl() in utils/assetUrl.js (important for split deploy).",
        ],
    )

    add_qa(
        doc,
        "Q13. How is an order placed?",
        [
            "Player browses GET /players/products (verified sellers only), builds a single-store cart.",
            "Checkout initiates Easypaisa: POST /players/orders/easypaisa/initiate → pending Payment.",
            "After payment proof: POST /players/orders → verifyEasypaisaPayment → finalizeProductOrder "
            "(atomic stock decrement + Order status paid) → notify business owner.",
            "Helpers: utils/productOrder.js, utils/atomicBooking.js; UI: PlayerShop.jsx / PlayerStore.jsx.",
        ],
    )

    # --- 8 Coach ---
    doc.add_heading("8. Coach Panel", level=1)

    add_qa(
        doc,
        "Q14. When a coach pays for subscription, how does the system know?",
        [
            "Stripe PaymentIntent is created (createCoachSubscriptionPaymentIntent) with metadata "
            "(purpose, action, userId) and amount in PKR.",
            "Frontend completes card payment (CoachSubscription.jsx + StripePaySection).",
            "Client calls POST /coaches/subscription with paymentIntentId.",
            "Server retrieves the PaymentIntent, verifies status succeeded, metadata, and amount.",
            "Creates Payment (type subscription, completed) and extends CoachProfile.platformSubscriptionRenewsAt "
            "by one month (extendCoachPlatformPeriod).",
            "requireCoachPlatformSubscription middleware blocks training features if renewsAt is expired "
            "(unless price is 0 / free).",
            "Files: coachController.js, utils/stripePayments.js, utils/coachPlatformSubscription.js, "
            "middleware/coachPlatformSubscription.js.",
        ],
    )

    add_qa(
        doc,
        "Q15. How does the coach receive training requests?",
        [
            "Player: POST /players/training-requests → TrainingRequest (pending) + notify coach.",
            "Coach: GET /coaches/training-requests lists requests.",
            "Accept/reject: PATCH /coaches/training-requests/:id (accept requires meeting schedule fields).",
            "Fees cleared + roll number: POST .../mark-fees-cleared.",
            "First session: POST .../start-session after fees cleared.",
            "UI: CoachRequests.jsx.",
        ],
    )

    add_qa(
        doc,
        "Q16. What is the evaluation logic?",
        [
            "Coach loads sport/category rubric (evaluationRubrics.js).",
            "Submits skill scores via POST /coaches/performance (addPerformance).",
            "normalizeSkillScores / deriveLegacyScores compute technique, fitness, attitude, overall.",
            "PerformanceEvaluation is stored; player is notified; data feeds weekly plans and recommendations.",
            "Coach must have a training relationship (session) with the player.",
            "UI: CoachPerformance.jsx; helpers: evaluationScores.js, skillGapAnalysis.js.",
        ],
    )

    add_qa(
        doc,
        "Q17. How is a training session created?",
        [
            "Core helper: scheduleTrainingSession() in coachController.js.",
            "Checks: accepted TrainingRequest, fees cleared, no schedule overlap, capacity.",
            "Creates TrainingSession (status scheduled) with scheduledAt, location, duration.",
            "First session from request: startTrainingFromRequest; later sessions: POST /coaches/training-sessions.",
            "Attendance can mark a session completed.",
        ],
    )

    # --- 9 Business ---
    doc.add_heading("9. Business Panel", level=1)

    add_qa(
        doc,
        "Q18. How are store setup images uploaded?",
        [
            "POST /business/store/logo and POST /business/store/banner with uploadImage.single('image').",
            "Controllers uploadStoreLogo / uploadStoreBanner save paths on BusinessProfile "
            "(storeLogoUrl, storeBannerUrl).",
            "UI: BusinessStoreSettings.jsx. Product images use POST /business/products/:id/images "
            "or generic /uploads/image.",
        ],
    )

    add_qa(
        doc,
        "Q19. When a user views a store, how does the system take them to that profile?",
        [
            "Equipment catalog: /player/shop → GET /players/products (includes store name/logo metadata).",
            "Product card links to /player/shop/store/:ownerId (PlayerStore.jsx).",
            "Page calls GET /players/stores/:ownerId (getBusinessStore) — verifies business_owner is verified "
            "and returns store branding + products.",
            "There is no separate public business profile page beyond the storefront view.",
        ],
    )

    add_qa(
        doc,
        "Q20. How does the business manage orders?",
        [
            "List: GET /business/orders (orders where businessOwner = current user).",
            "Update: PATCH /business/orders/:id with status (processing, shipped, completed, …) "
            "and optional trackingNumber; player is notified.",
            "Orders arrive as paid from player Easypaisa checkout (finalizeProductOrder).",
            "UI: BusinessOrders.jsx; optional sales report via GET /business/reports/sales.",
        ],
    )

    # --- 10 ---
    doc.add_heading("10. Key Source Files Reference", level=1)
    add_table(
        doc,
        ["Area", "Primary files"],
        [
            ("Auth & registration", "backend/controllers/authController.js, routes/auth.routes.js, middleware/auth.js"),
            ("Admin verification", "backend/controllers/adminController.js, routes/admin.routes.js"),
            ("Documents / uploads", "middleware/upload.js, models/VerificationDocument.js, uploadController.js"),
            ("Recommendations", "playerController.js (getRecommendations), services/aiCoachEngine.js"),
            ("Coach subscription", "coachController.js, utils/stripePayments.js, coachPlatformSubscription.js"),
            ("Training lifecycle", "coachController.js, models/TrainingRequest.js, TrainingSession.js"),
            ("Evaluations", "evaluationRubrics.js, evaluationScores.js, CoachPerformance.jsx"),
            ("Shop & orders", "playerController.js, productOrder.js, atomicBooking.js, businessController.js"),
            ("Frontend auth", "frontend/src/hooks/useAuth.jsx, services/api.js"),
            ("Deploy", "render.yaml, frontend/vercel.json, ecosystem.config.js"),
        ],
    )

    doc.add_paragraph()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    er = end.add_run("— End of Document —")
    set_run_font(er, 10, True, RGBColor(0x88, 0x88, 0x88))

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Sports Ecosystem Platform · FYP26-CS-G22 · Confidential for academic use")
    set_run_font(fr, 9, False, RGBColor(0x99, 0x99, 0x99))

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
